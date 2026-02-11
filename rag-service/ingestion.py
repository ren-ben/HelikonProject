import io
import uuid
from datetime import datetime, timezone

from langchain_core.documents import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter

from vector_store import get_vectorstore

_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)


# Text extraction — page-aware for PDFs

def extract_pages(file_bytes: bytes, filename: str) -> list[tuple[int | None, str]]:
    """Return list of (page_number, text) tuples.

    PDFs get real 1-based page numbers; DOCX/TXT get page_number=None.
    """
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

    if ext == "pdf":
        import PyPDF2
        reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
        pages = []
        for i, page in enumerate(reader.pages):
            text = page.extract_text() or ""
            if text.strip():
                pages.append((i + 1, text))
        return pages if pages else [(None, "")]

    if ext == "docx":
        from docx import Document as DocxDocument
        doc = DocxDocument(io.BytesIO(file_bytes))
        return [(None, "\n".join(p.text for p in doc.paragraphs))]

    if ext == "txt":
        return [(None, file_bytes.decode("utf-8"))]

    raise ValueError(f"Unsupported file type: .{ext}  (allowed: .pdf, .docx, .txt)")


# Ingest pipeline

def ingest_document(
    file_bytes: bytes,
    filename: str,
    user_id: str,
    metadata: dict | None = None,
) -> dict:
    """Extract, chunk, embed, store. Returns {doc_id, chunk_count, filename}."""
    pages = extract_pages(file_bytes, filename)

    doc_id = str(uuid.uuid4())
    now = datetime.now(timezone.utc).isoformat()

    # chroma metadata values must be str, int, float, bool.
    base_meta: dict = {
        "user_id": user_id,
        "doc_id": doc_id,
        "filename": filename,
        "uploaded_at": now,
    }
    if metadata:
        for key, val in metadata.items():
            base_meta[key] = ",".join(val) if isinstance(val, list) else val

    documents = []
    chunk_idx = 0
    for page_num, page_text in pages:
        page_chunks = _splitter.split_text(page_text)
        for chunk in page_chunks:
            meta = {**base_meta, "chunk_index": chunk_idx}
            if page_num is not None:
                meta["page_number"] = page_num
            documents.append(Document(page_content=chunk, metadata=meta))
            chunk_idx += 1

    get_vectorstore().add_documents(documents)

    return {"doc_id": doc_id, "chunk_count": len(documents), "filename": filename}
